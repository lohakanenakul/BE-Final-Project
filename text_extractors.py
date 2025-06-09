import PyPDF2
import pdfplumber
from docx import Document
import logging
import io

class TextExtractor:
    """Base class for text extraction"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def extract_text(self, file_path):
        """Extract text from file - to be implemented by subclasses"""
        raise NotImplementedError

class PDFExtractor(TextExtractor):
    """Extract text from PDF files"""
    
    def extract_text(self, file_path):
        """
        Extract text from PDF file using multiple methods for better accuracy
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text content
        """
        text = ""
        
        try:
            # First try with pdfplumber (better for complex layouts)
            text = self._extract_with_pdfplumber(file_path)
            
            if not text or len(text.strip()) < 100:
                # Fallback to PyPDF2
                text = self._extract_with_pypdf2(file_path)
            
        except Exception as e:
            self.logger.error(f"Error extracting PDF text: {str(e)}")
            
        return text.strip()
    
    def _extract_with_pdfplumber(self, file_path):
        """Extract text using pdfplumber"""
        text = ""
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            self.logger.warning(f"pdfplumber extraction failed: {str(e)}")
            
        return text
    
    def _extract_with_pypdf2(self, file_path):
        """Extract text using PyPDF2 as fallback"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        
        except Exception as e:
            self.logger.warning(f"PyPDF2 extraction failed: {str(e)}")
            
        return text

class DOCXExtractor(TextExtractor):
    """Extract text from DOCX files"""
    
    def extract_text(self, file_path):
        """
        Extract text from DOCX file
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text content
        """
        text = ""
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            # Extract text from headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {str(e)}")
            
        return text.strip()
    
    def extract_formatting_info(self, file_path):
        """
        Extract additional formatting information that might be useful
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            dict: Formatting information
        """
        formatting_info = {
            'bold_text': [],
            'italic_text': [],
            'headings': []
        }
        
        try:
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                # Check for headings
                if paragraph.style.name.startswith('Heading'):
                    formatting_info['headings'].append({
                        'level': paragraph.style.name,
                        'text': paragraph.text
                    })
                
                # Check for bold and italic text
                for run in paragraph.runs:
                    if run.bold and run.text.strip():
                        formatting_info['bold_text'].append(run.text.strip())
                    if run.italic and run.text.strip():
                        formatting_info['italic_text'].append(run.text.strip())
                        
        except Exception as e:
            self.logger.warning(f"Error extracting formatting info: {str(e)}")
            
        return formatting_info
